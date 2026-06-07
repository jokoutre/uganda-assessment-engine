import streamlit as st
import pandas as pd
from google import genai
from fpdf import FPDF
from docx import Document
import io
import os
import re

# 1. Web Framework Interface Setting
st.set_page_config(page_title="Uganda CBC Worksheet Engine", layout="wide")

st.title("🇺🇬 NCDC Competency-Based Assessment Engine (V2.1)")
st.write("Constructing high-end PDFs, editable Word docs, and Assessment Reports with Pattern Intelligence.")
st.divider()

# 2. Control Sidebar Deck
st.sidebar.header("Worksheet Controls")
api_key = st.sidebar.text_input("Enter Gemini API Key:", type="password")

class_level = st.sidebar.selectbox("Target Class:", ["Primary 5", "Primary 6", "Primary 7"])
subject = st.sidebar.selectbox("Core Subject:", [
    "Mathematics", "Integrated Science", "Social Studies (SST)", 
    "English Language", "Kiswahili", "Religious Education (RE)"
])
term = st.sidebar.selectbox("Academic Term:", ["Term I", "Term II", "Term III"])
specific_topic = st.sidebar.text_input("Curriculum Topic:", placeholder="e.g., Punctuation, Geometry, Malaria Prevention")

ability_level = st.sidebar.selectbox("Learner Ability Level:", [
    "Average Ability (Standard CBC)",
    "Lower Ability (Foundational/Guided)",
    "High Ability (Challenging/Analytical)"
])

assignment_type = st.sidebar.radio("Assignment Type:", ["Classwork", "Homework"])
question_count = st.sidebar.slider("Number of Questions:", min_value=1, max_value=15, value=5)

st.sidebar.subheader("Curriculum Reference")
uploaded_file = st.sidebar.file_uploader("Upload NCDC Reference PDF:", type=["pdf"])

# --- ADVANCED DESIGN BLUEPRINT MATRIX PUBLISHER (PDF) ---
def create_blueprint_pdf(class_name, subj_name, term_name, topic_name, assign_type, raw_ai_text):
    replacements = {
        '\u2013': '-', '\u2014': '-', '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"', '\u2026': '...', '**': '' 
    }
    for bad_char, good_char in replacements.items():
        raw_ai_text = raw_ai_text.replace(bad_char, good_char)
        
    cleaned_lines = []
    for line in raw_ai_text.split('\n'):
        stripped = line.strip()
        if stripped in ['*', '-', '•', '>', '·']: continue
        cleaned_lines.append(line)
    
    raw_ai_text = '\n'.join(cleaned_lines)
    raw_ai_text = re.sub(r'\n{3,}', '\n\n', raw_ai_text)

    pdf = FPDF(orientation='P', unit='mm', format='A4')
    pdf.add_page()
    pdf.set_margins(left=25.4, top=25.4, right=25.4)
    pdf.set_auto_page_break(auto=True, margin=25.4) 
    
    text_color_rgb = (38, 38, 38)      
    line_color_rgb = (166, 166, 166)    
    fill_color_rgb = (242, 242, 242)    
    border_grey_rgb = (127, 127, 127)   
    
    pdf.set_text_color(*text_color_rgb)
    pdf.set_draw_color(*text_color_rgb)
    
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_x(25.4)
    pdf.cell(159.2, 10, "UGANDA PRIMARY SCHOOL ASSESSMENT ENGINE", ln=True, align="C")
    
    pdf.set_font("Helvetica", "B", 12)
    pdf.set_x(25.4)
    pdf.cell(159.2, 8, f"{assign_type.upper()} EVALUATION SHEET: {topic_name.upper()}", ln=True, align="C")
    pdf.ln(4)
    
    pdf.set_font("Helvetica", "", 10)
    pdf.set_x(25.4)
    pdf.cell(100, 8, "LEARNER'S NAME: __________________________________", border=0)
    pdf.cell(59.2, 8, f"CLASS: {class_name}", border=0, ln=True)
    pdf.set_x(25.4)
    pdf.cell(100, 8, "DATE: ____________________________________________", border=0)
    pdf.cell(59.2, 8, f"SUBJECT: {subj_name} ({term_name})", border=0, ln=True)
    
    pdf.set_line_width(0.6)
    pdf.line(25.4, pdf.get_y() + 4, 184.6, pdf.get_y() + 4)
    pdf.ln(8)
    
    in_scenario = False
    
    lines = raw_ai_text.split("\n")
    for line in lines:
        clean_line = line.strip()
        if not clean_line: continue
            
        if "[SCENARIO_START]" in clean_line:
            in_scenario = True
            pdf.set_draw_color(*border_grey_rgb)
            pdf.set_fill_color(*fill_color_rgb)
            pdf.set_line_width(0.35) 
            if pdf.get_y() > 210: pdf.add_page()
            scenario_y_start = pdf.get_y()
            pdf.set_font("Helvetica", "I", 11) 
            continue
            
        elif "[SCENARIO_END]" in clean_line:
            in_scenario = False
            scenario_y_end = pdf.get_y()
            box_height = scenario_y_end - scenario_y_start + 4
            pdf.rect(25.4, scenario_y_start - 2, 159.2, box_height, style="D")
            pdf.set_draw_color(*text_color_rgb)
            pdf.ln(4)
            continue
            
        if "[DIAGRAM_PLACEHOLDER:" in clean_line:
            match = re.search(r'\[DIAGRAM_PLACEHOLDER:\s*(.*?)\]', clean_line)
            diagram_desc = match.group(1).upper() if match else "RELEVANT DIAGRAM"
            if pdf.get_y() > 230: pdf.add_page()
            
            pdf.ln(2)
            pdf.set_draw_color(*line_color_rgb)
            pdf.set_line_width(0.5)
            current_y = pdf.get_y()
            
            pdf.rect(25.4, current_y, 159.2, 45)
            pdf.set_y(current_y + 20)
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(*line_color_rgb)
            pdf.set_x(25.4)
            pdf.cell(159.2, 6, f"[ ILLUSTRATION SPACE: PASTE {diagram_desc} HERE ]", align="C")
            
            pdf.set_text_color(*text_color_rgb)
            pdf.set_draw_color(*text_color_rgb)
            pdf.set_y(current_y + 50)
            continue
            
        if "[SPACE_RESERVED]" in clean_line:
            line_text = clean_line.replace("[SPACE_RESERVED]", "")
            if line_text:
                pdf.set_font("Helvetica", "", 12)
                safe_text = line_text.encode('latin-1', 'replace').decode('latin-1')
                pdf.set_x(25.4)
                pdf.multi_cell(159.2, 6, safe_text)
                
            current_y = pdf.get_y()
            
            if subj_name == "Mathematics":
                box_height_mm = 63.5 
                if current_y + box_height_mm > 270:
                    pdf.add_page()
                    current_y = pdf.get_y()
                pdf.set_draw_color(*border_grey_rgb)
                pdf.set_line_width(0.53) 
                pdf.rect(25.4, current_y + 2, 159.2, box_height_mm)
                pdf.set_y(current_y + box_height_mm - 6)
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_x(25.4)
                pdf.cell(30, 6, "  Answer Box Space: ", border=0)
                pdf.set_y(current_y + box_height_mm + 4)
            else:
                pdf.ln(4)
                pdf.set_draw_color(*line_color_rgb)
                pdf.set_line_width(0.25)
                for _ in range(3):
                    if pdf.get_y() > 270: pdf.add_page()
                    pdf.line(25.4, pdf.get_y() + 6, 184.6, pdf.get_y() + 6)
                    pdf.ln(10)
                pdf.ln(2)
            pdf.set_draw_color(*text_color_rgb)
            
        else:
            safe_clean_line = clean_line.encode('latin-1', 'replace').decode('latin-1')
            pdf.set_x(25.4)
            if in_scenario:
                pdf.set_font("Helvetica", "I", 11)
                pdf.multi_cell(159.2, 6, safe_clean_line)
            else:
                pdf.set_font("Helvetica", "", 12) 
                pdf.multi_cell(159.2, 6, safe_clean_line)
                pdf.ln(1.5)

    if assign_type == "Homework":
        if pdf.get_y() > 220: pdf.add_page()
        pdf.ln(6)
        pdf.set_draw_color(*text_color_rgb)
        pdf.set_line_width(0.5)
        box_top = pdf.get_y()
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_x(25.4)
        pdf.cell(159.2, 6, "  PARENT'S / GUARDIAN'S VERIFICATION BOX", ln=True)
        pdf.set_font("Helvetica", "", 10)
        pdf.set_x(25.4)
        pdf.cell(159.2, 6, "  Dear Parent/Guardian, please check your child's homework assignment metrics and tick:", ln=True)
        pdf.set_x(25.4)
        pdf.cell(159.2, 6, "  [  ] My child completed this assessment completely independently.", ln=True)
        pdf.set_x(25.4)
        pdf.cell(159.2, 6, "  [  ] My child required structural guidance to process the Part B scenario.", ln=True)
        pdf.ln(3)
        pdf.set_x(25.4)
        pdf.cell(159.2, 6, "  Parent's Name: ___________________________          Signature: ___________________________", ln=True)
        pdf.ln(2)
        box_bottom = pdf.get_y()
        pdf.rect(25.4, box_top, 159.2, box_bottom - box_top)
        
    return pdf.output()

# --- MS WORD DOCUMENT ENGINE ---
def create_word_doc(class_name, subj_name, term_name, topic_name, assign_type, raw_ai_text):
    doc = Document()
    
    doc.add_heading('UGANDA PRIMARY SCHOOL ASSESSMENT ENGINE', 0)
    doc.add_heading(f'{assign_type.upper()} EVALUATION: {topic_name.upper()}', 1)
    
    p = doc.add_paragraph()
    p.add_run(f"LEARNER'S NAME: ____________________\t\tCLASS: {class_name}\n").bold = True
    p.add_run(f"DATE: ______________________________\t\tSUBJECT: {subj_name} ({term_name})\n").bold = True
    
    word_text = raw_ai_text.replace('[SCENARIO_START]', '\n--- SCENARIO CONTEXT START ---\n')
    word_text = word_text.replace('[SCENARIO_END]', '\n--- SCENARIO CONTEXT END ---\n')
    word_text = word_text.replace('[SPACE_RESERVED]', '\n\n[   Teacher: Leave working space here   ]\n\n')
    word_text = word_text.replace('**', '')
    
    word_text = re.sub(
        r'\[DIAGRAM_PLACEHOLDER:\s*(.*?)\]', 
        r'\n\n[ 🖼️ TEACHER: Paste diagram of \1 here ]\n\n', 
        word_text
    )
    
    doc.add_paragraph(word_text)
    
    if assign_type == "Homework":
        doc.add_heading("PARENT'S / GUARDIAN'S VERIFICATION", 2)
        doc.add_paragraph("[  ] My child completed this assessment completely independently.")
        doc.add_paragraph("[  ] My child required structural guidance to process the Part B application scenario.")
        doc.add_paragraph("Parent's Name: _____________________    Signature: _____________________")
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 4. Main Infrastructure Automation Tower
if st.sidebar.button("Generate Master Assessment Documents ✨"):
    if not api_key:
        st.sidebar.error("Please supply a valid Google AI Studio API Key.")
    elif not specific_topic:
        st.sidebar.error("Please explicitly declare a curriculum lesson topic.")
    else:
        with st.spinner(f"Compiling V2.1 {ability_level} materials, mapping competencies, & generating report..."):
            try:
                client = genai.Client(api_key=api_key)
                ai_contents = []
                
                if uploaded_file is not None:
                    with open("temp_manifest.pdf", "wb") as f:
                        f.write(uploaded_file.getbuffer())
                    uploaded_doc = client.files.upload(file="temp_manifest.pdf")
                    ai_contents.append(uploaded_doc)
                
                # --- V2 DYNAMIC BLUEPRINT INJECTION WITH PATTERN INTELLIGENCE ---
                blueprint_rules = ""
                
                if subject == "English Language":
                    blueprint_rules = f"""
                    ASSESSMENT PATTERN INTELLIGENCE: English is NOT a comprehension-only subject. 
                    Intelligently select the most appropriate format based on the topic: '{specific_topic}'.
                    Valid Formats: Grammar (Tenses, punctuation, concord), Vocabulary (Synonyms, antonyms), Functional Writing (Letters, reports), Dialogue Completion, Situational Reading (Notices, ads), Information Interpretation (Timetables, charts), or Comprehension.
                    
                    CRITICAL RULE: IF AND ONLY IF the generated task is specifically a Reading Comprehension activity, you MUST activate the English Comprehension Blueprint and use this exact distribution:
                    """
                    if "Lower" in ability_level:
                        blueprint_rules += "40% Literal/Recall, 20% Vocabulary, 20% Inferential, 10% Cause and Effect, 10% Main Idea."
                    elif "Average" in ability_level:
                        blueprint_rules += "30% Literal/Recall, 20% Vocabulary, 20% Inferential, 10% Cause and Effect, 10% Main Idea, 10% Evaluation."
                    else:
                        blueprint_rules += "20% Literal/Recall, 20% Vocabulary, 30% Inferential, 10% Main Idea, 10% Prediction, 10% Evaluation."
                        
                elif subject == "Mathematics":
                    blueprint_rules = """
                    ASSESSMENT PATTERN INTELLIGENCE: Balance questions appropriately across valid math formats based on the topic:
                    - Mechanical Skills (Addition, subtraction, fractions, decimals)
                    - Financial Literacy (Market transactions, shopping, budgeting, profit/loss)
                    - Data Interpretation (Tables, bar graphs, pie charts, records)
                    - Geometry (Shapes, area, perimeter, volume, spatial reasoning)
                    - Real-Life Problem Solving (Farming, school management, construction)
                    """
                elif subject == "Integrated Science":
                    blueprint_rules = """
                    ASSESSMENT PATTERN INTELLIGENCE: Balance questions appropriately across valid science formats based on the topic:
                    - Scientific Knowledge
                    - Apparatus Interpretation (Charcoal stoves, beehives, farm tools)
                    - Diagram Analysis (Labeling, identifying mistakes, explaining functions)
                    - Disease Prevention Scenarios (Malaria, hygiene, nutrition)
                    - Agricultural Problem Solving (Soil erosion, pest control, drought)
                    - Home and Community Application (Kitchen science, waste management, water purification)
                    """
                elif subject == "Social Studies (SST)":
                    blueprint_rules = """
                    ASSESSMENT PATTERN INTELLIGENCE: Balance questions appropriately across valid SST formats based on the topic:
                    - Knowledge Questions
                    - Sketch Map Interpretation (Uganda maps, village maps)
                    - Civic Responsibility Scenarios (Broken bridge, dirty water source, sanitation)
                    - Economic Decision Making (Markets, roads, mobile money)
                    - Environmental Management (Wetlands, forests, soil conservation)
                    - Citizenship and Governance (Local councils, national symbols)
                    """
                elif subject == "Kiswahili":
                    blueprint_rules = """
                    ASSESSMENT PATTERN INTELLIGENCE: Balance questions appropriately across valid Kiswahili formats based on the topic:
                    - Vocabulary Development
                    - Translation Activities (English <-> Kiswahili)
                    - Conversational Situations (Greetings, market conversations)
                    - Sentence Construction & Grammar
                    - Short Reading Comprehension (Only when appropriate, focusing on basic retrieval: Nani, Lini, Wapi, Nini).
                    """
                elif subject == "Religious Education (RE)":
                    blueprint_rules = """
                    ASSESSMENT PATTERN INTELLIGENCE: Balance questions appropriately across valid RE formats based on the topic:
                    - Knowledge & Scripture/Teaching Interpretation
                    - Moral Dilemmas (Returning lost money, honesty, bullying, respect)
                    - Value Application (Forgiveness, hard work, compassion, stewardship)
                    - Community and School-Based Ethics (Leadership, responsibility, service)
                    """

                system_prompt = f"""
                You are a senior curriculum executive designing competency-based materials for Uganda.
                
                *** PHASE 1: LOGICAL VALIDATION GATE ***
                Verify that the topic '{specific_topic}' makes logical sense for the subject '{subject}'. 
                CRITICAL EXCEPTION: For 'English Language' and 'Kiswahili', ALMOST ANY topic is valid because they serve as themes for dialogues, vocabulary, and grammar. Do NOT flag topics as contradictory for language subjects.
                If a topic is blatantly contradictory for a non-language subject, output EXACTLY this string and NOTHING ELSE:
                ===VALIDATION_ERROR===: The topic '{specific_topic}' does not logically align with the subject '{subject}'.
                If you output the validation error, YOU MUST STOP IMMEDIATELY. Do not generate the worksheet.
                
                *** PHASE 2: GENERATION RULES & DIFFERENTIATION ***
                Generate content strictly tailored for {class_level}, {term}, {subject}.
                DIFFERENTIATION CALIBRATION ("{ability_level}"):
                - Lower Ability: Basic vocabulary/short sentences. Simple recall/guided steps.
                - Average Ability: Standard grade-level language and standard reasoning steps.
                - High Ability: Advanced vocabulary/complex sentences. Multi-step logic/abstract analysis.
                
                1. REVENUE CONTEXT: Ugandan Shillings (UGX) must NEVER display decimals.
                2. CBC SCENARIO GENERATOR (PART B): All Part B questions MUST use authentic Ugandan community scenarios encouraging critical thinking. Wrap scenarios in '[SCENARIO_START]' and '[SCENARIO_END]'.
                3. RESPONSE CHANNELS: Every question must end with '[SPACE_RESERVED]'. If a question requires a long response, output '[SPACE_RESERVED]' multiple times on consecutive lines.
                
                *** PHASE 3: SUBJECT BLUEPRINTS & ASSESSMENT INTELLIGENCE ***
                {blueprint_rules}
                
                COMPETENCY MAPPING: Internally tag every generated question in the Answer Key with its CBC competency (e.g., Critical Thinking, Communication, Ethics). Hide these tags from the Student Worksheet.
                BALANCE CHECKER: Audit your questions. Ensure a mix of Knowledge, Understanding, Application, Reasoning, and Problem Solving. Automatically rebalance if biased towards pure recall.
                
                *** PHASE 4: THE MULTI-SPLIT OUTPUT PROTOCOL ***
                You MUST output three distinct sections separated by EXACT tokens:
                [Student Worksheet Text]
                ===ANSWER_KEY_SPLIT===
                [Teacher Answer Key with Competency Tags]
                ===TEACHER_REPORT_SPLIT===
                [Teacher Assessment Report - Include Summary, Question Distribution, Competencies Covered, and AI Recommendations]

                *** PHASE 5: UNIVERSAL DIAGRAM INTELLIGENCE ***
                If a visual or sketch map is required, output: [DIAGRAM_PLACEHOLDER: brief description of image]. Do NOT draw text art.
                
                *** PHASE 6: TYPOGRAPHY & LAYOUT STRICT RULES ***
                1. Write in standard sentence case. DO NOT write full sentences in ALL CAPS.
                2. DO NOT use markdown formatting like **bold** or *italics*.
                3. NEVER use bullet points (asterisks * or dashes -).
                4. For main questions, use standard numbers: 1. 2. 3.
                5. For sub-questions, use letters inline: a) b) c).
                """
                
                user_prompt = f"Create a {question_count}-question evaluation assignment on: '{specific_topic}' calibrated for {ability_level}."
                ai_contents.append(user_prompt)
                
                response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=ai_contents, 
                    config={"system_instruction": system_prompt}
                )
                
                raw_output = response.text.strip()
                
                if raw_output.startswith("===VALIDATION_ERROR==="):
                    error_message = raw_output.replace("===VALIDATION_ERROR===:", "").strip()
                    st.error(f"🚨 Logic Conflict Detected: {error_message}")
                    
                else:
                    student_text = raw_output
                    teacher_text = "Answer Key missing or formatting error."
                    report_text = "Teacher Assessment Report missing or formatting error."
                    
                    if "===ANSWER_KEY_SPLIT===" in student_text:
                        parts = student_text.split("===ANSWER_KEY_SPLIT===", 1)
                        student_text = parts[0]
                        remainder = parts[1]
                        
                        if "===TEACHER_REPORT_SPLIT===" in remainder:
                            teacher_text, report_text = remainder.split("===TEACHER_REPORT_SPLIT===", 1)
                        else:
                            teacher_text = remainder
                    
                    pdf_binary = create_blueprint_pdf(class_level, subject, term, specific_topic, assignment_type, student_text)
                    word_binary = create_word_doc(class_level, subject, term, specific_topic, assignment_type, student_text)
                    
                    tab1, tab2, tab3 = st.tabs(["📄 Export Assessment File", "🔑 Teacher Matrix Key", "📊 Assessment Report"])
                    
                    with tab1:
                        st.success(f"Your CBC {ability_level} assessment is audited, compiled, and ready!")
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.download_button(
                                label="📥 Download Locked PDF (For Printing)",
                                data=bytes(pdf_binary),
                                file_name=f"{class_level}_{subject.replace(' ', '_')}_{specific_topic.replace(' ', '_')}.pdf",
                                mime="application/pdf"
                            )
                        with col2:
                            st.download_button(
                                label="📝 Download Editable Word Doc",
                                data=word_binary,
                                file_name=f"{class_level}_{subject.replace(' ', '_')}_{specific_topic.replace(' ', '_')}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        st.write("---")
                        st.write("### Raw Text Contents Review:")
                        st.write(
                            student_text
                            .replace("[SPACE_RESERVED]", "")
                            .replace("[SCENARIO_START]", "")
                            .replace("[SCENARIO_END]", "")
                        )
                        
                    with tab2:
                        st.download_button(
                            label="💾 Download Teacher Answer Key (.txt)",
                            data=teacher_text,
                            file_name=f"{class_level}_{subject}_Answers.txt"
                        )
                        st.write("### Competency-Mapped Answer Key:")
                        st.write(teacher_text)
                        
                    with tab3:
                        st.download_button(
                            label="📈 Download AI Assessment Report (.txt)",
                            data=report_text,
                            file_name=f"{class_level}_{subject}_Report.txt"
                        )
                        st.write("### Internal CBC Assessment Audit:")
                        st.write(report_text)
                        
                if uploaded_file is not None:
                    client.files.delete(name=uploaded_doc.name)
                    
            except Exception as e:
                st.error(f"System compiling error occurred: {e}")